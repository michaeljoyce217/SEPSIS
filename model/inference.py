# model/inference.py
# Sepsis Appeal Engine - Appeal Letter Generation
#
# GENERATION PIPELINE (reads prepared data from case tables):
# 1. Load knowledge base (gold letters, Propel definitions)
# 2. Read case data from tables (written by featurization_inference.py)
# 3. Vector search → Find best matching gold letter
# 4. Generate appeal letter → Using prepared clinical data
# 5. Assess strength → Evaluate against criteria
# 6. Export to DOCX → With conflicts appendix if applicable
#
# PREREQUISITE: Run featurization_inference.py first to prepare case data.
#
# Run on Databricks Runtime 15.4 LTS ML

# =============================================================================
# CELL 1: Configuration
# =============================================================================
import os
import math
import json
import re
from datetime import datetime
from pyspark.sql import SparkSession

spark = SparkSession.builder.getOrCreate()

# -----------------------------------------------------------------------------
# Processing Configuration
# -----------------------------------------------------------------------------
SCOPE_FILTER = "sepsis"
SEPSIS_DRG_CODES = ["870", "871", "872"]
MATCH_SCORE_THRESHOLD = 0.7

# Default template path
DEFAULT_TEMPLATE_PATH = "/Workspace/Repos/mijo8881@mercy.net/fudgesicle/utils/gold_standard_appeals_sepsis_only/default_sepsis_appeal_template.docx"

# Output configuration
EXPORT_TO_DOCX = True
DOCX_OUTPUT_BASE = "/Workspace/Repos/mijo8881@mercy.net/fudgesicle/utils/outputs"

# -----------------------------------------------------------------------------
# Environment Setup
# -----------------------------------------------------------------------------
# NOTE: Data lives in prod catalog, but we write to our environment's catalog.
# This is intentional - we query from prod but can only write to our own env.
trgt_cat = os.environ.get('trgt_cat', 'dev')
spark.sql('USE CATALOG prod;')

# Knowledge base tables
GOLD_LETTERS_TABLE = f"{trgt_cat}.fin_ds.fudgesicle_gold_letters"
PROPEL_DATA_TABLE = f"{trgt_cat}.fin_ds.fudgesicle_propel_data"

# Case data tables (written by featurization_inference.py)
CASE_DENIAL_TABLE = f"{trgt_cat}.fin_ds.fudgesicle_case_denial"
CASE_CLINICAL_TABLE = f"{trgt_cat}.fin_ds.fudgesicle_case_clinical"
CASE_STRUCTURED_SUMMARY_TABLE = f"{trgt_cat}.fin_ds.fudgesicle_case_structured_summary"
CASE_CONFLICTS_TABLE = f"{trgt_cat}.fin_ds.fudgesicle_case_conflicts"

print(f"Catalog: {trgt_cat}")

# =============================================================================
# CELL 2: Azure Credentials and Client
# =============================================================================
from openai import AzureOpenAI

AZURE_OPENAI_KEY = dbutils.secrets.get(scope='idp_etl', key='az-openai-key1')
AZURE_OPENAI_ENDPOINT = dbutils.secrets.get(scope='idp_etl', key='az-openai-base')

openai_client = AzureOpenAI(
    api_key=AZURE_OPENAI_KEY,
    azure_endpoint=AZURE_OPENAI_ENDPOINT,
    api_version="2024-10-21"
)

print("Azure OpenAI client initialized")

# =============================================================================
# CELL 3: Load Knowledge Base
# =============================================================================
print("\n" + "="*60)
print("LOADING KNOWLEDGE BASE")
print("="*60)

# Load gold letters
print("\nLoading gold standard letters...")
gold_letters_cache = []
try:
    gold_letters_df = spark.sql(f"""
        SELECT letter_id, source_file, payor, denial_text, rebuttal_text, denial_embedding, metadata
        FROM {GOLD_LETTERS_TABLE}
    """)
    gold_letters = gold_letters_df.collect()

    gold_letters_cache = [
        {
            "letter_id": row["letter_id"],
            "source_file": row["source_file"],
            "payor": row["payor"],
            "denial_text": row["denial_text"],
            "appeal_text": row["rebuttal_text"],
            "denial_embedding": list(row["denial_embedding"]) if row["denial_embedding"] else None,
            "metadata": dict(row["metadata"]) if row["metadata"] else {},
        }
        for row in gold_letters
    ]
    print(f"  Loaded {len(gold_letters_cache)} gold standard letters")
except Exception as e:
    print(f"  Warning: Could not load gold letters: {e}")

# Load default template
print("\nLoading default template...")
default_template = None
try:
    if os.path.exists(DEFAULT_TEMPLATE_PATH):
        from docx import Document as DocxDocument
        doc = DocxDocument(DEFAULT_TEMPLATE_PATH)
        template_text = "\n\n".join([para.text for para in doc.paragraphs if para.text.strip()])

        default_template = {
            "letter_id": "default_template",
            "source_file": os.path.basename(DEFAULT_TEMPLATE_PATH),
            "payor": "Generic",
            "denial_text": None,
            "appeal_text": template_text,
            "denial_embedding": None,
            "metadata": {"is_default_template": "true"},
        }
        print(f"  Loaded default template: {len(template_text)} chars")
    else:
        print(f"  Warning: Default template not found at {DEFAULT_TEMPLATE_PATH}")
except Exception as e:
    print(f"  Warning: Could not load default template: {e}")

# Load Propel definitions
print("\nLoading Propel clinical definitions...")
propel_definitions = {}
try:
    propel_df = spark.sql(f"""
        SELECT condition_name, definition_summary, definition_text
        FROM {PROPEL_DATA_TABLE}
    """)
    for row in propel_df.collect():
        definition = row["definition_summary"] or row["definition_text"]
        propel_definitions[row["condition_name"]] = definition
        print(f"  {row['condition_name']}: {len(definition)} chars")
except Exception as e:
    print(f"  Warning: Could not load propel definitions: {e}")

# =============================================================================
# CELL 4: Load Case Data from Tables
# =============================================================================
print("\n" + "="*60)
print("LOADING CASE DATA FROM TABLES")
print("="*60)

def load_case_data():
    """Load all case data from tables written by featurization_inference.py."""
    case_data = {}

    # Load denial data
    print("\nLoading denial data...")
    try:
        denial_row = spark.sql(f"SELECT * FROM {CASE_DENIAL_TABLE}").collect()
        if denial_row:
            row = denial_row[0]
            case_data["account_id"] = row["account_id"]
            case_data["denial_text"] = row["denial_text"]
            case_data["denial_embedding"] = list(row["denial_embedding"]) if row["denial_embedding"] else None
            case_data["payor"] = row["payor"]
            case_data["original_drg"] = row["original_drg"]
            case_data["proposed_drg"] = row["proposed_drg"]
            case_data["is_sepsis"] = row["is_sepsis"]
            print(f"  Account: {case_data['account_id']}")
            print(f"  Payor: {case_data['payor']}")
            print(f"  Sepsis: {case_data['is_sepsis']}")
        else:
            print("  ERROR: No denial data found. Run featurization_inference.py first.")
            return None
    except Exception as e:
        print(f"  ERROR: Could not load denial data: {e}")
        return None

    # Load clinical data
    print("\nLoading clinical data...")
    try:
        clinical_row = spark.sql(f"SELECT * FROM {CASE_CLINICAL_TABLE}").collect()
        if clinical_row:
            row = clinical_row[0]
            case_data["patient_name"] = row["patient_name"]
            case_data["patient_dob"] = row["patient_dob"]
            case_data["facility_name"] = row["facility_name"]
            case_data["date_of_service"] = row["date_of_service"]
            case_data["extracted_notes"] = json.loads(row["extracted_notes"]) if row["extracted_notes"] else {}
            print(f"  Patient: {case_data['patient_name']}")
            print(f"  Notes loaded: {len([v for v in case_data['extracted_notes'].values() if v != 'Not available'])}")
        else:
            print("  WARNING: No clinical data found")
            case_data["patient_name"] = "Unknown"
            case_data["patient_dob"] = ""
            case_data["facility_name"] = "Mercy Hospital"
            case_data["date_of_service"] = ""
            case_data["extracted_notes"] = {}
    except Exception as e:
        print(f"  WARNING: Could not load clinical data: {e}")
        case_data["extracted_notes"] = {}

    # Load structured summary
    print("\nLoading structured data summary...")
    try:
        structured_row = spark.sql(f"SELECT * FROM {CASE_STRUCTURED_SUMMARY_TABLE}").collect()
        if structured_row:
            case_data["structured_summary"] = structured_row[0]["structured_summary"]
            print(f"  Summary: {len(case_data['structured_summary'])} chars")
        else:
            print("  WARNING: No structured summary found")
            case_data["structured_summary"] = "No structured data available."
    except Exception as e:
        print(f"  WARNING: Could not load structured summary: {e}")
        case_data["structured_summary"] = "No structured data available."

    # Load conflicts
    print("\nLoading conflicts...")
    try:
        conflicts_row = spark.sql(f"SELECT * FROM {CASE_CONFLICTS_TABLE}").collect()
        if conflicts_row:
            row = conflicts_row[0]
            conflicts_list = json.loads(row["conflicts"]) if row["conflicts"] else []
            case_data["conflicts"] = {
                "conflicts": conflicts_list,
                "recommendation": row["recommendation"] or ""
            }
            print(f"  Conflicts: {len(conflicts_list)}")
        else:
            print("  No conflicts data found")
            case_data["conflicts"] = {"conflicts": [], "recommendation": ""}
    except Exception as e:
        print(f"  WARNING: Could not load conflicts: {e}")
        case_data["conflicts"] = {"conflicts": [], "recommendation": ""}

    return case_data


# Load case data
case_data = load_case_data()

if not case_data:
    print("\n" + "="*60)
    print("ERROR: Case data not found. Run featurization_inference.py first.")
    print("="*60)
else:
    # =============================================================================
    # CELL 5: Vector Search Function
    # =============================================================================

    def find_best_gold_letter(denial_embedding):
        """Find the best matching gold letter using cosine similarity."""
        if not gold_letters_cache or not denial_embedding:
            return None, 0.0

        best_score = 0.0
        best_match = None

        for letter in gold_letters_cache:
            if letter["denial_embedding"]:
                vec1 = denial_embedding
                vec2 = letter["denial_embedding"]
                dot_product = sum(a * b for a, b in zip(vec1, vec2))
                norm1 = math.sqrt(sum(a * a for a in vec1))
                norm2 = math.sqrt(sum(b * b for b in vec2))
                if norm1 > 0 and norm2 > 0:
                    similarity = dot_product / (norm1 * norm2)
                    if similarity > best_score:
                        best_score = similarity
                        best_match = letter

        if best_match and best_score >= MATCH_SCORE_THRESHOLD:
            return best_match, best_score
        elif default_template:
            return default_template, 0.0
        else:
            return None, 0.0

    # =============================================================================
    # CELL 6: Writer Prompt
    # =============================================================================
    WRITER_PROMPT = '''You are a clinical coding expert writing a DRG validation appeal letter for Mercy Hospital.

# Original Denial Letter
{denial_letter_text}

# Clinical Notes (PRIMARY EVIDENCE - from physician documentation)
## Discharge Summary
{discharge_summary}

## H&P Note
{hp_note}

## Progress Notes
{progress_note}

## Consult Notes
{consult_note}

## ED Notes
{ed_notes}

## Initial Assessments
{initial_assessment}

## ED Triage Notes
{ed_triage}

## ED Provider Notes
{ed_provider_note}

## Addendum Note
{addendum_note}

## Hospital Course
{hospital_course}

## Subjective & Objective
{subjective_objective}

## Assessment & Plan Note
{assessment_plan}

## Nursing Note
{nursing_note}

## Code Documentation
{code_documentation}

# Structured Data Summary (SUPPORTING EVIDENCE - from labs, vitals, meds)
{structured_data_summary}

# Official Clinical Definition
{clinical_definition_section}

# Gold Standard Letter
{gold_letter_section}

# Patient Information
Name: {patient_name}
DOB: {patient_dob}
Hospital Account #: {hsp_account_id}
Date of Service: {date_of_service}
Facility: {facility_name}
Original DRG: {original_drg}
Proposed DRG: {proposed_drg}
Payor: {payor}

# Instructions
{gold_letter_instructions}
1. READ THE DENIAL LETTER - extract the payor address, reviewer name, claim numbers
2. ADDRESS EACH DENIAL ARGUMENT - quote the payer, then refute
3. CITE CLINICAL EVIDENCE from provider notes FIRST, then structured data
4. INCLUDE TIMESTAMPS with clinical values
5. QUANTIFY ORGAN DYSFUNCTION using SOFA criteria when available:
   - Reference specific values: lactate, MAP, creatinine, platelets, bilirubin, GCS, PaO2/FiO2
   - Example: "Patient demonstrated cardiovascular dysfunction with MAP of 63 requiring vasopressor support"
6. Follow the Mercy Hospital template structure exactly

Return ONLY the letter text.'''

    # =============================================================================
    # CELL 7: Assessment Functions
    # =============================================================================

    ASSESSMENT_PROMPT = '''You are evaluating the strength of a sepsis DRG appeal letter.

═══ PROPEL SEPSIS CRITERIA ═══
{propel_definition}

═══ DENIAL LETTER ═══
{denial_text}

═══ GOLD LETTER TEMPLATE USED ═══
{gold_letter_text}

═══ AVAILABLE CLINICAL EVIDENCE ═══
{extracted_clinical_data}

═══ STRUCTURED DATA SUMMARY ═══
{structured_summary}

═══ GENERATED APPEAL LETTER ═══
{generated_letter}

═══ EVALUATION INSTRUCTIONS ═══
Evaluate this appeal letter and provide:

1. OVERALL SCORE (1-10) and RATING (LOW for 1-4, MODERATE for 5-7, HIGH for 8-10)
2. SUMMARY (2-3 sentences)
3. DETAILED BREAKDOWN with scores and specific findings

Return ONLY valid JSON in this format:
{{
  "overall_score": <1-10>,
  "overall_rating": "<LOW|MODERATE|HIGH>",
  "summary": "<2-3 sentence summary>",
  "propel_criteria": {{
    "score": <1-10>,
    "findings": [
      {{"status": "<present|could_strengthen|missing>", "item": "<description>"}}
    ]
  }},
  "argument_structure": {{
    "score": <1-10>,
    "findings": [
      {{"status": "<present|could_strengthen|missing>", "item": "<description>"}}
    ]
  }},
  "evidence_quality": {{
    "clinical_notes": {{"score": <1-10>, "findings": [...]}},
    "structured_data": {{"score": <1-10>, "findings": [...]}}
  }}
}}'''


    def assess_appeal_strength(generated_letter, propel_definition, denial_text,
                               extracted_notes, gold_letter_text, structured_summary):
        """Assess the strength of a generated appeal letter."""
        print("  Running strength assessment...")

        # Format extracted notes
        notes_summary = []
        for note_type, content in extracted_notes.items():
            if content and content != "Not available":
                truncated = content[:2000] + "..." if len(content) > 2000 else content
                notes_summary.append(f"## {note_type}\n{truncated}")
        extracted_clinical_data = "\n\n".join(notes_summary) if notes_summary else "No clinical notes available"

        try:
            response = openai_client.chat.completions.create(
                model="gpt-4.1",
                messages=[
                    {"role": "system", "content": "You are a clinical appeal quality assessor. Return only valid JSON."},
                    {"role": "user", "content": ASSESSMENT_PROMPT.format(
                        propel_definition=propel_definition or "Propel criteria not available",
                        denial_text=denial_text[:5000] if denial_text else "Denial text not available",
                        gold_letter_text=gold_letter_text[:3000] if gold_letter_text else "No gold letter template used",
                        extracted_clinical_data=extracted_clinical_data,
                        structured_summary=structured_summary[:3000] if structured_summary else "No structured data",
                        generated_letter=generated_letter
                    )}
                ],
                temperature=0,
                max_tokens=2000
            )

            raw_response = response.choices[0].message.content.strip()

            # Parse JSON
            if raw_response.startswith("```"):
                raw_response = raw_response.split("```")[1]
                if raw_response.startswith("json"):
                    raw_response = raw_response[4:]
                raw_response = raw_response.strip()

            assessment = json.loads(raw_response)

            if "overall_score" in assessment:
                assessment["overall_score"] = max(1, min(10, int(assessment["overall_score"])))

            print(f"  Assessment complete: {assessment.get('overall_score', '?')}/10 - {assessment.get('overall_rating', '?')}")
            return assessment

        except json.JSONDecodeError as e:
            print(f"  Warning: Assessment JSON parse error: {e}")
            return None
        except Exception as e:
            print(f"  Warning: Assessment failed: {e}")
            return None


    def format_assessment_for_docx(assessment):
        """Format assessment dict into text for DOCX output."""
        if not assessment:
            return "Assessment unavailable\n\nPlease review letter manually."

        status_symbols = {"present": "✓", "could_strengthen": "△", "missing": "✗"}
        lines = []

        lines.append(f"Overall Strength: {assessment.get('overall_score', '?')}/10 - {assessment.get('overall_rating', '?')}")
        lines.append("")
        lines.append(f"Summary: {assessment.get('summary', 'No summary available')}")
        lines.append("")
        lines.append("Detailed Breakdown:")
        lines.append("─" * 55)

        # Propel Criteria
        propel = assessment.get("propel_criteria", {})
        lines.append(f"PROPEL CRITERIA: {propel.get('score', '?')}/10")
        for finding in propel.get("findings", []):
            symbol = status_symbols.get(finding.get("status", ""), "?")
            lines.append(f"  {symbol} {finding.get('item', '')[:50]}")

        # Argument Structure
        argument = assessment.get("argument_structure", {})
        lines.append(f"\nARGUMENT STRUCTURE: {argument.get('score', '?')}/10")
        for finding in argument.get("findings", []):
            symbol = status_symbols.get(finding.get("status", ""), "?")
            lines.append(f"  {symbol} {finding.get('item', '')[:50]}")

        # Evidence Quality
        evidence = assessment.get("evidence_quality", {})
        clinical = evidence.get("clinical_notes", {})
        structured = evidence.get("structured_data", {})
        lines.append(f"\nEVIDENCE - Clinical Notes: {clinical.get('score', '?')}/10")
        for finding in clinical.get("findings", []):
            symbol = status_symbols.get(finding.get("status", ""), "?")
            lines.append(f"  {symbol} {finding.get('item', '')[:48]}")
        lines.append(f"\nEVIDENCE - Structured Data: {structured.get('score', '?')}/10")
        for finding in structured.get("findings", []):
            symbol = status_symbols.get(finding.get("status", ""), "?")
            lines.append(f"  {symbol} {finding.get('item', '')[:48]}")

        lines.append("─" * 55)
        return "\n".join(lines)

    # =============================================================================
    # CELL 8: Main Processing Pipeline
    # =============================================================================
    print("\n" + "="*60)
    print("INFERENCE - APPEAL LETTER GENERATION")
    print("="*60)

    account_id = case_data["account_id"]
    extracted_notes = case_data["extracted_notes"]
    structured_summary = case_data["structured_summary"]
    conflicts_result = case_data["conflicts"]

    # -------------------------------------------------------------------------
    # STEP 1: Find best gold letter match
    # -------------------------------------------------------------------------
    print("\nStep 1: Finding best gold letter match...")
    gold_letter, gold_letter_score = find_best_gold_letter(case_data["denial_embedding"])
    if gold_letter:
        is_default = gold_letter.get("metadata", {}).get("is_default_template") == "true"
        if is_default:
            print(f"  Match: {gold_letter['source_file']} (score: N/A - default template)")
        else:
            print(f"  Match: {gold_letter['source_file']} (score: {gold_letter_score:.3f})")
    else:
        print("  No match found")

    # -------------------------------------------------------------------------
    # STEP 2: Generate appeal letter
    # -------------------------------------------------------------------------
    print("\nStep 2: Generating appeal letter...")

    # Build gold letter section
    if gold_letter:
        is_default = gold_letter.get("metadata", {}).get("is_default_template") == "true"
        if is_default:
            gold_letter_section = f"## APPEAL TEMPLATE\n{gold_letter['appeal_text']}"
            gold_letter_instructions = "**NOTE: Using default template as structural guide.**\n"
        else:
            gold_letter_section = f"## WINNING APPEAL (Score: {gold_letter_score:.3f})\nPayor: {gold_letter.get('payor')}\n\n{gold_letter['appeal_text']}"
            gold_letter_instructions = "**CRITICAL: Learn from this winning appeal.**\n"
    else:
        gold_letter_section = "No template available."
        gold_letter_instructions = ""

    # Clinical definition
    if case_data["is_sepsis"] and "sepsis" in propel_definitions:
        clinical_definition_section = f"## OFFICIAL SEPSIS DEFINITION\n{propel_definitions['sepsis']}"
    else:
        clinical_definition_section = "No specific definition loaded."

    # Build prompt
    writer_prompt = WRITER_PROMPT.format(
        denial_letter_text=case_data["denial_text"],
        discharge_summary=extracted_notes.get("discharge_summary", "Not available"),
        hp_note=extracted_notes.get("hp_note", "Not available"),
        progress_note=extracted_notes.get("progress_note", "Not available"),
        consult_note=extracted_notes.get("consult_note", "Not available"),
        ed_notes=extracted_notes.get("ed_notes", "Not available"),
        initial_assessment=extracted_notes.get("initial_assessment", "Not available"),
        ed_triage=extracted_notes.get("ed_triage", "Not available"),
        ed_provider_note=extracted_notes.get("ed_provider_note", "Not available"),
        addendum_note=extracted_notes.get("addendum_note", "Not available"),
        hospital_course=extracted_notes.get("hospital_course", "Not available"),
        subjective_objective=extracted_notes.get("subjective_objective", "Not available"),
        assessment_plan=extracted_notes.get("assessment_plan", "Not available"),
        nursing_note=extracted_notes.get("nursing_note", "Not available"),
        code_documentation=extracted_notes.get("code_documentation", "Not available"),
        structured_data_summary=structured_summary,
        clinical_definition_section=clinical_definition_section,
        gold_letter_section=gold_letter_section,
        gold_letter_instructions=gold_letter_instructions,
        patient_name=case_data.get("patient_name", ""),
        patient_dob=case_data.get("patient_dob", ""),
        hsp_account_id=account_id,
        date_of_service=case_data.get("date_of_service", ""),
        facility_name=case_data.get("facility_name", "Mercy Hospital"),
        original_drg=case_data.get("original_drg") or "Unknown",
        proposed_drg=case_data.get("proposed_drg") or "Unknown",
        payor=case_data.get("payor", "Unknown"),
    )

    # Generate letter
    writer_response = openai_client.chat.completions.create(
        model="gpt-4.1",
        messages=[
            {"role": "system", "content": "You are a clinical coding expert writing DRG appeal letters."},
            {"role": "user", "content": writer_prompt}
        ],
        temperature=0.2,
        max_tokens=4000
    )

    letter_text = writer_response.choices[0].message.content.strip()
    print(f"  Generated {len(letter_text)} character letter")

    # -------------------------------------------------------------------------
    # STEP 3: Assess appeal strength
    # -------------------------------------------------------------------------
    print("\nStep 3: Assessing appeal strength...")
    propel_def = propel_definitions.get("sepsis") if case_data["is_sepsis"] else None
    gold_text = gold_letter.get("appeal_text", "") if gold_letter else ""

    assessment = assess_appeal_strength(
        letter_text, propel_def, case_data["denial_text"],
        extracted_notes, gold_text, structured_summary
    )

    # -------------------------------------------------------------------------
    # STEP 4: Export to DOCX
    # -------------------------------------------------------------------------
    if EXPORT_TO_DOCX:
        print("\nStep 4: Exporting to DOCX...")
        from docx import Document
        from docx.shared import Pt, RGBColor

        def add_markdown_paragraph(doc, text):
            """Add paragraph with markdown bold converted to Word bold."""
            p = doc.add_paragraph()
            parts = re.split(r'(\*\*.*?\*\*)', text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)
            return p

        os.makedirs(DOCX_OUTPUT_BASE, exist_ok=True)

        doc = Document()
        doc.add_heading('Appeal Letter', level=1)

        # Metadata
        meta = doc.add_paragraph()
        meta.add_run("Generated: ").bold = True
        meta.add_run(f"{datetime.now().strftime('%Y-%m-%d %H:%M')}\n")
        meta.add_run("Patient: ").bold = True
        meta.add_run(f"{case_data.get('patient_name', 'Unknown')}\n")
        meta.add_run("Payor: ").bold = True
        meta.add_run(f"{case_data.get('payor', 'Unknown')}\n")
        meta.add_run("Gold Letter: ").bold = True
        is_default = gold_letter.get("metadata", {}).get("is_default_template") == "true" if gold_letter else False
        if is_default:
            meta.add_run(f"{gold_letter['source_file'] if gold_letter else 'None'} (score: N/A)\n")
        else:
            meta.add_run(f"{gold_letter['source_file'] if gold_letter else 'None'} (score: {gold_letter_score:.3f})\n")

        # Assessment section
        doc.add_paragraph("═" * 55)
        header = doc.add_paragraph()
        header.add_run("APPEAL STRENGTH ASSESSMENT (Internal Review Only)").bold = True
        doc.add_paragraph("═" * 55)

        for line in format_assessment_for_docx(assessment).split('\n'):
            p = doc.add_paragraph(line)
            p.paragraph_format.space_after = Pt(0)

        doc.add_paragraph("═" * 55)

        # Conflicts appendix (if any)
        if conflicts_result.get('conflicts'):
            doc.add_paragraph()
            conflict_header = doc.add_paragraph()
            conflict_header.add_run("⚠️ CONFLICTS DETECTED - REQUIRES CDI REVIEW").bold = True
            doc.add_paragraph("─" * 55)
            for conflict in conflicts_result['conflicts']:
                p = doc.add_paragraph(conflict)
                p.paragraph_format.space_after = Pt(4)
            if conflicts_result.get('recommendation'):
                doc.add_paragraph()
                rec = doc.add_paragraph()
                rec.add_run("Recommendation: ").bold = True
                rec.add_run(conflicts_result['recommendation'])
            doc.add_paragraph("─" * 55)

        doc.add_paragraph()

        # Letter content
        for paragraph in letter_text.split('\n\n'):
            if paragraph.strip():
                p = add_markdown_paragraph(doc, paragraph.strip())
                p.paragraph_format.space_after = Pt(12)

        # Save
        patient_name = case_data.get('patient_name', 'Unknown')
        safe_name = "".join(c for c in patient_name if c.isalnum() or c in (' ', '-', '_')).strip()
        filename = f"{account_id}_{safe_name}_appeal.docx"
        filepath = os.path.join(DOCX_OUTPUT_BASE, filename)
        doc.save(filepath)

        print(f"  Saved: {filepath}")

    # -------------------------------------------------------------------------
    # SUMMARY
    # -------------------------------------------------------------------------
    print("\n" + "="*60)
    print("INFERENCE COMPLETE")
    print("="*60)
    print(f"Patient: {case_data.get('patient_name', 'Unknown')}")
    print(f"Account: {account_id}")
    print(f"Letter length: {len(letter_text)} chars")
    if assessment:
        print(f"Strength: {assessment.get('overall_score', '?')}/10 - {assessment.get('overall_rating', '?')}")
    if conflicts_result.get('conflicts'):
        print(f"Conflicts: {len(conflicts_result['conflicts'])} (see appendix in DOCX)")
    if EXPORT_TO_DOCX:
        print(f"Output: {filepath}")

print("\nInference complete.")
